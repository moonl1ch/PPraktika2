from docx import Document
from db.database import Database

class WordExporter:
    def __init__(self, db):
        self.db = db

    def export_services(self, filename='services.docx'):
        services = self.db.get_services()
        doc = Document()
        doc.add_heading('Услуги', 0)

        for service in services:
            doc.add_paragraph(f"{service[1]} - {service[2]} руб.")

        doc.save(filename)

    def export_employees(self, filename='employees.docx'):
        employees = self.db.get_employees()
        doc = Document()
        doc.add_heading('Сотрудники', 0)

        for employee in employees:
            doc.add_paragraph(f"{employee[1]} - {employee[2]}")

        doc.save(filename)

    def export_clients(self, filename='clients.docx'):
        clients = self.db.get_clients()
        doc = Document()
        doc.add_heading('Клиенты', 0)

        for client in clients:
            doc.add_paragraph(f"{client[1]} - {client[2]}")

        doc.save(filename)

    def export_items(self, filename='items.docx'):
        items = self.db.get_items()
        doc = Document()
        doc.add_heading('Вещи', 0)

        for item in items:
            doc.add_paragraph(f"{item[1]} (Клиент ID: {item[2]}, Услуга ID: {item[3]})")

        doc.save(filename)

    def export_all(self, filename='all_data.docx'):
        services = self.db.get_services()
        employees = self.db.get_employees()
        clients = self.db.get_clients()
        items = self.db.get_items()

        doc = Document()
        doc.add_heading('Все данные', 0)

        doc.add_heading('Услуги', level=1)
        for service in services:
            doc.add_paragraph(f"{service[1]} - {service[2]} руб.")

        doc.add_heading('Сотрудники', level=1)
        for employee in employees:
            doc.add_paragraph(f"{employee[1]} - {employee[2]}")

        doc.add_heading('Клиенты', level=1)
        for client in clients:
            doc.add_paragraph(f"{client[1]} - {client[2]}")

        doc.add_heading('Вещи', level=1)
        for item in items:
            doc.add_paragraph(f"{item[1]} (Клиент ID: {item[2]}, Услуга ID: {item[3]})")

        doc.save(filename)
